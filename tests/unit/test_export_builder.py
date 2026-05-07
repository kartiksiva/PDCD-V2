"""Unit tests for export_builder — evidence bundle, PDF, Markdown, DOCX (PRD §8.10)."""

from __future__ import annotations

import base64
import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../../backend"))

import pytest
from app.export_builder import (
    _classify_anchor_type,
    _derive_roles,
    build_evidence_bundle,
    build_export_docx,
    build_export_markdown,
    build_export_pdf,
)

TINY_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7+6wAAAABJRU5ErkJggg=="
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _step(step_id: str, summary: str, anchors: list | None = None) -> dict:
    return {
        "id": step_id,
        "summary": summary,
        "actor": "Analyst",
        "source_anchors": anchors or [],
    }


def _sipoc_row(process_step: str, source_anchor: str, step_anchor: list | None = None) -> dict:
    return {
        "supplier": "Client",
        "input": "Request",
        "process_step": process_step,
        "output": "Result",
        "customer": "Team",
        "source_anchor": source_anchor,
        "step_anchor": step_anchor or [],
    }


DRAFT_WITH_ANCHORS = {
    "pdd": {
        "purpose": "Test process",
        "scope": "Unit test scope",
        "steps": [
            _step("step-1", "First step", [{"anchor": "00:00:00-00:00:10", "confidence": 0.9}]),
            _step("step-2", "Second step", [{"anchor": "00:01:15", "confidence": 0.75}]),
        ],
    },
    "sipoc": [
        _sipoc_row("Log request", "00:00:00-00:00:10", ["step-1"]),
        _sipoc_row("Validate", "00:01:15", ["step-2"]),
    ],
}

DRAFT_EMPTY = {}

JOB_WITH_SIGNALS = {
    "agent_signals": {"evidence_strength": "high"},
    "input_manifest": {
        "video": {
            "frame_extraction_policy": {"sample_interval_sec": 5, "ocr_enabled": True}
        }
    },
    "extracted_evidence": {
        "evidence_items": [
            {"anchor": "00:00:00-00:00:10", "ocr_text": "Screen shows login form"},
        ]
    },
}

JOB_MINIMAL = {}


# ---------------------------------------------------------------------------
# _classify_anchor_type
# ---------------------------------------------------------------------------

class TestClassifyAnchorType:
    def test_timestamp_range(self):
        assert _classify_anchor_type("00:00:00-00:00:10") == "timestamp_range"

    def test_timestamp_point(self):
        assert _classify_anchor_type("00:01:15") == "timestamp_range"

    def test_frame_id(self):
        assert _classify_anchor_type("frame-42") == "frame_id"
        assert _classify_anchor_type("FRAME-001") == "frame_id"

    def test_section_label(self):
        assert _classify_anchor_type("intro-section") == "section_label"
        assert _classify_anchor_type("Step A overview") == "section_label"

    def test_empty_string(self):
        assert _classify_anchor_type("") == "missing"


# ---------------------------------------------------------------------------
# build_evidence_bundle
# ---------------------------------------------------------------------------

class TestBuildEvidenceBundle:
    def test_collects_pdd_step_anchors(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)
        values = [a["anchor_value"] for a in bundle["linked_anchors"]]
        assert "00:00:00-00:00:10" in values
        assert "00:01:15" in values

    def test_linked_step_ids_populated(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        by_val = {a["anchor_value"]: a for a in bundle["linked_anchors"]}
        assert "step-1" in by_val["00:00:00-00:00:10"]["linked_step_ids"]
        assert "step-2" in by_val["00:01:15"]["linked_step_ids"]

    def test_linked_sipoc_rows_populated(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        by_val = {a["anchor_value"]: a for a in bundle["linked_anchors"]}
        assert 0 in by_val["00:00:00-00:00:10"]["linked_sipoc_rows"]
        assert 1 in by_val["00:01:15"]["linked_sipoc_rows"]

    def test_anchor_type_classification(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        by_val = {a["anchor_value"]: a for a in bundle["linked_anchors"]}
        assert by_val["00:00:00-00:00:10"]["anchor_type"] == "timestamp_range"
        assert by_val["00:01:15"]["anchor_type"] == "timestamp_range"

    def test_ocr_snippet_attached(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)
        by_val = {a["anchor_value"]: a for a in bundle["linked_anchors"]}
        assert by_val["00:00:00-00:00:10"]["ocr_snippet"] == "Screen shows login form"

    def test_ocr_snippet_none_when_no_match(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)
        by_val = {a["anchor_value"]: a for a in bundle["linked_anchors"]}
        assert by_val["00:01:15"]["ocr_snippet"] is None

    def test_evidence_strength_from_job_signals(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)
        assert bundle["evidence_strength"] == "high"

    def test_evidence_strength_none_when_absent(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        assert bundle["evidence_strength"] is None

    def test_total_linked_anchors_count(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        assert bundle["total_linked_anchors"] == len(bundle["linked_anchors"])

    def test_frame_policy_from_job(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)
        assert bundle["frame_policy"]["sample_interval_sec"] == 5

    def test_frame_policy_none_when_absent(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        assert bundle["frame_policy"] is None

    def test_prd_810_only_linked_anchors_included(self):
        """Anchors not referenced by any step or SIPOC row must not appear."""
        draft_no_sipoc_anchor = {
            "pdd": {
                "purpose": "p", "scope": "s",
                "steps": [_step("step-1", "x", [{"anchor": "00:00:05", "confidence": 0.8}])],
            },
            "sipoc": [_sipoc_row("y", "", ["step-1"])],  # source_anchor empty
        }
        bundle = build_evidence_bundle(draft_no_sipoc_anchor, JOB_MINIMAL)
        # "00:00:05" is linked to step-1, so it should be included
        values = [a["anchor_value"] for a in bundle["linked_anchors"]]
        assert "00:00:05" in values
        # empty sipoc source_anchor must not create a phantom entry
        assert "" not in values

    def test_empty_draft_returns_empty_bundle(self):
        bundle = build_evidence_bundle(DRAFT_EMPTY, JOB_MINIMAL)
        assert bundle["linked_anchors"] == []
        assert bundle["total_linked_anchors"] == 0

    def test_step_anchor_deduplication(self):
        """Same anchor value in two steps merges into one entry with both step IDs."""
        draft = {
            "pdd": {
                "purpose": "p", "scope": "s",
                "steps": [
                    _step("step-1", "a", [{"anchor": "00:00:30", "confidence": 0.8}]),
                    _step("step-2", "b", [{"anchor": "00:00:30", "confidence": 0.7}]),
                ],
            },
            "sipoc": [],
        }
        bundle = build_evidence_bundle(draft, JOB_MINIMAL)
        assert len(bundle["linked_anchors"]) == 1
        entry = bundle["linked_anchors"][0]
        assert set(entry["linked_step_ids"]) == {"step-1", "step-2"}

    def test_frame_captures_note_present(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        assert "frame_captures_note" in bundle
        assert bundle["frame_captures_note"]

    def test_anchor_id_unique(self):
        bundle = build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_MINIMAL)
        ids = [a["anchor_id"] for a in bundle["linked_anchors"]]
        assert len(ids) == len(set(ids))

    def test_sipoc_only_anchor_included(self):
        """Anchor present only in SIPOC (not in any PDD step) must still appear."""
        draft = {
            "pdd": {"purpose": "p", "scope": "s", "steps": []},
            "sipoc": [_sipoc_row("Do something", "00:05:00", [])],
        }
        bundle = build_evidence_bundle(draft, JOB_MINIMAL)
        values = [a["anchor_value"] for a in bundle["linked_anchors"]]
        assert "00:05:00" in values

    def test_frame_captures_reject_unsafe_storage_keys(self):
        draft = {
            "pdd": {
                "purpose": "p",
                "scope": "s",
                "steps": [_step("step-1", "x", [{"anchor": "00:00:10", "confidence": 0.8}])],
            },
            "sipoc": [_sipoc_row("x", "00:00:10", ["step-1"])],
        }
        job = {
            "extracted_evidence": {
                "evidence_items": [
                    {
                        "metadata": {
                            "frame_storage_keys": [
                                ("job-123/frames/frame_0001.jpg", 10.0),
                                ("/tmp/frame.jpg", 11.0),
                                ("../frames/frame_0002.jpg", 12.0),
                                ("./frames/frame_0003.jpg", 13.0),
                            ]
                        }
                    }
                ]
            },
            "agent_signals": {
                "frame_storage_keys": [
                    ("job-123/frames/frame_0001.jpg", 10.0),
                    ("..\\frames\\frame_0004.jpg", 14.0),
                ]
            },
        }
        bundle = build_evidence_bundle(draft, job)
        captures = bundle["frame_captures"]
        assert captures == [{"storage_key": "job-123/frames/frame_0001.jpg", "timestamp_sec": 10.0, "linked_anchor_ids": ["a-1"]}]


# ---------------------------------------------------------------------------
# build_export_markdown
# ---------------------------------------------------------------------------

class TestBuildExportMarkdown:
    def _bundle(self):
        return build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)

    def test_contains_pdd_title(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "# Process Definition Document" in md

    def test_contains_purpose_and_scope(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "Test process" in md
        assert "Unit test scope" in md

    def test_contains_steps_with_anchors(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "step-1" in md
        assert "00:00:00-00:00:10" in md

    def test_contains_sipoc_section(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "### 2.6 SIPOC (Supplier / Input / Process / Output / Customer)" in md
        assert "Log request" in md

    def test_contains_evidence_bundle_section(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "### Evidence Bundle Manifest" in md

    def test_evidence_strength_shown(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "high" in md

    def test_anchor_table_present(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "| Anchor |" in md

    def test_ocr_snippet_in_table(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "Screen shows login form" in md

    def test_frame_captures_note_present(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "Frame captures" in md

    def test_exceptions_table_uses_trigger_field(self):
        draft = {
            "pdd": {
                "purpose": "p",
                "scope": "s",
                "steps": [],
                "exceptions": [
                    {
                        "scenario": "Validation failed",
                        "trigger": "Missing mandatory field",
                        "action_required": "Request correction",
                        "owner": "Analyst",
                    }
                ],
            },
            "sipoc": [],
        }
        md = build_export_markdown(draft, build_evidence_bundle(draft, JOB_MINIMAL))
        assert "Missing mandatory field" in md

    def test_raci_uses_approval_matrix_for_non_actor_roles(self):
        draft = {
            "pdd": {
                "purpose": "p",
                "scope": "s",
                "steps": [
                    {
                        "id": "step-01",
                        "summary": "Capture request",
                        "actor": "Analyst",
                        "source_anchors": [{"anchor": "00:00:01-00:00:05", "confidence": 0.9}],
                    }
                ],
                "roles": ["Analyst", "Manager", "Compliance"],
                "exceptions": [],
            },
            "approval_matrix": [
                {"role": "Manager", "responsibility": "A"},
                {"role": "Compliance", "responsibility": "C"},
            ],
            "sipoc": [],
        }
        md = build_export_markdown(draft, build_evidence_bundle(draft, JOB_MINIMAL))
        assert "| Capture request | R | A | C |" in md

    def test_empty_draft_returns_fallback(self):
        md = build_export_markdown({}, build_evidence_bundle({}, {}))
        assert "No finalized draft available" in md

    def test_semantic_review_section_hidden_when_empty(self):
        md = build_export_markdown(DRAFT_WITH_ANCHORS, self._bundle())
        assert "## 5A. AI Semantic Review Notes" not in md

    def test_semantic_review_section_rendered_when_flags_present(self):
        bundle = self._bundle()
        bundle["llm_semantic_flags"] = [
            {
                "code": "coverage_gap",
                "message": "Evidence ev-03 is not mapped to any PDD step.",
                "evidence_id": "ev-03",
                "step_id": "",
                "anchor": "00:00:20-00:00:30",
                "snippet": "Analyst reconciles ERP billing",
            }
        ]
        md = build_export_markdown(DRAFT_WITH_ANCHORS, bundle)
        assert "## 5A. AI Semantic Review Notes" in md
        assert "ev-03" in md
        assert "00:00:20-00:00:30" in md

    def test_no_anchors_shows_fallback_message(self):
        draft = {"pdd": {"purpose": "p", "scope": "s", "steps": []}, "sipoc": []}
        bundle = {"linked_anchors": [], "evidence_strength": None, "frame_captures_note": ""}
        md = build_export_markdown(draft, bundle)
        assert "No linked evidence anchors found" in md


# ---------------------------------------------------------------------------
# build_export_pdf
# ---------------------------------------------------------------------------

class TestBuildExportPdf:
    def _bundle(self):
        return build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)

    def test_returns_bytes(self):
        result = build_export_pdf(DRAFT_WITH_ANCHORS, self._bundle())
        assert isinstance(result, bytes)

    def test_pdf_magic_bytes(self):
        result = build_export_pdf(DRAFT_WITH_ANCHORS, self._bundle())
        assert result[:4] == b"%PDF"

    def test_empty_draft_does_not_crash(self):
        bundle = build_evidence_bundle({}, {})
        result = build_export_pdf({}, bundle)
        assert isinstance(result, bytes)
        assert result[:4] == b"%PDF"

    def test_no_linked_anchors_does_not_crash(self):
        bundle = {"linked_anchors": [], "evidence_strength": "low", "frame_captures_note": ""}
        result = build_export_pdf(DRAFT_WITH_ANCHORS, bundle)
        assert isinstance(result, bytes)

    def test_pdf_embeds_frame_from_bytes_map(self):
        bundle = {
            "linked_anchors": [],
            "evidence_strength": "high",
            "frame_captures_note": "",
            "frame_captures": [{"storage_key": "job-123/frames/frame_0001.jpg", "timestamp_sec": 12.0}],
        }
        result = build_export_pdf(
            DRAFT_WITH_ANCHORS,
            bundle,
            frame_bytes_map={"job-123/frames/frame_0001.jpg": TINY_PNG_BYTES},
        )
        assert isinstance(result, bytes)
        assert result[:4] == b"%PDF"

    def test_pdf_skips_missing_frame_gracefully(self):
        bundle = {
            "linked_anchors": [],
            "evidence_strength": "high",
            "frame_captures_note": "",
            "frame_captures": [{"storage_key": "job-123/frames/missing.jpg", "timestamp_sec": 12.0}],
        }
        result = build_export_pdf(DRAFT_WITH_ANCHORS, bundle, frame_bytes_map={})
        assert isinstance(result, bytes)
        assert result[:4] == b"%PDF"

    def test_pdf_contains_section_headings(self):
        """PDF should include major structural section headings."""
        from pypdf import PdfReader
        import io as _io

        draft = {
            "pdd": {
                "purpose": "Test purpose",
                "scope": "Test scope",
                "steps": [_step("step-1", "First step", [])],
                "exceptions": [{"scenario": "Timeout", "trigger": "No response", "action": "Escalate", "owner": "Team"}],
                "process_controls": [{"control_id": "C-01", "description": "Log all events", "owner": "Ops", "frequency": "Always", "tool": "SIEM"}],
            },
            "sipoc": [_sipoc_row("Do something", "00:00:00", ["step-1"])],
            "approval_matrix": [{"role": "Manager", "responsibility": "A"}],
            "automation_opportunities": [{"description": "Automate X", "quantification": "1h", "automation_signal": "high"}],
            "faqs": [{"question": "What is this?", "answer": "A process."}],
        }
        bundle = {"linked_anchors": [], "evidence_strength": "medium", "frame_captures_note": ""}
        pdf_bytes = build_export_pdf(draft, bundle)
        reader = PdfReader(_io.BytesIO(pdf_bytes))
        full_text = "\n".join(page.extract_text() or "" for page in reader.pages)

        assert "Process Steps" in full_text
        assert "Process Exceptions" in full_text
        assert "Approval Matrix" in full_text
        assert "Automation" in full_text

    def test_pdf_exceptions_uses_trigger_field(self):
        """PDF should render the trigger field, not leave it as 'Needs Review'."""
        from pypdf import PdfReader
        import io as _io

        draft = {
            "pdd": {
                "purpose": "P",
                "scope": "S",
                "steps": [],
                "exceptions": [{"scenario": "Scenario A", "trigger": "User cancels", "action": "Stop", "owner": "Lead"}],
            },
            "sipoc": [],
        }
        bundle = {"linked_anchors": [], "evidence_strength": "low", "frame_captures_note": ""}
        pdf_bytes = build_export_pdf(draft, bundle)
        reader = PdfReader(_io.BytesIO(pdf_bytes))
        full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "User cancels" in full_text

    def test_pdf_sipoc_renders_supplier_column(self):
        """PDF should render SIPOC table with column headers."""
        from pypdf import PdfReader
        import io as _io

        draft = {
            "pdd": {"purpose": "P", "scope": "S", "steps": []},
            "sipoc": [_sipoc_row("Step A", "00:00:00", [])],
        }
        bundle = {"linked_anchors": [], "evidence_strength": "low", "frame_captures_note": ""}
        pdf_bytes = build_export_pdf(draft, bundle)
        reader = PdfReader(_io.BytesIO(pdf_bytes))
        full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "Supplier" in full_text

    def test_pdf_em_dash_does_not_raise(self):
        """PDF builder must not crash when draft contains em dash characters."""
        draft = {
            "pdd": {
                "purpose": "Purpose \u2014 with em dash",
                "scope": "Scope",
                "steps": [{"id": "s-1", "summary": "Step \u2014 detail", "actor": "Analyst", "source_anchors": []}],
            },
            "sipoc": [],
        }
        bundle = {"linked_anchors": [], "evidence_strength": "low", "frame_captures_note": ""}
        result = build_export_pdf(draft, bundle)
        assert result[:4] == b"%PDF"


# ---------------------------------------------------------------------------
# _derive_roles
# ---------------------------------------------------------------------------

class TestDeriveRoles:
    def _draft_with_roles(self, roles, step_actors=None):
        steps = [{"id": "s-1", "summary": "Do it", "actor": a, "source_anchors": []} for a in (step_actors or [])]
        return {"pdd": {"roles": roles, "steps": steps}}

    def test_keeps_clean_roles(self):
        draft = self._draft_with_roles(["Intake Analyst", "Call Team", "Management"])
        result = _derive_roles(draft)
        assert "Intake Analyst" in result
        assert "Call Team" in result
        assert "Management" in result

    def test_excludes_slash_composites(self):
        draft = self._draft_with_roles(["Analyst", "Customer / Call Agent", "System / Analyst"])
        result = _derive_roles(draft)
        assert "Analyst" in result
        assert "Customer / Call Agent" not in result
        assert "System / Analyst" not in result

    def test_excludes_session_meta_labels_case_insensitive(self):
        draft = self._draft_with_roles(["Analyst", "Service Provider", "customer", "CUSTOMER"])
        result = _derive_roles(draft)
        assert "Analyst" in result
        assert "Service Provider" not in result
        assert "customer" not in result
        assert "CUSTOMER" not in result

    def test_slash_composite_in_step_actor_excluded(self):
        draft = self._draft_with_roles([], step_actors=["Analyst", "Customer / Agent"])
        result = _derive_roles(draft)
        assert "Analyst" in result
        assert "Customer / Agent" not in result

    def test_deduplicates_roles(self):
        draft = self._draft_with_roles(["Analyst", "Analyst"], step_actors=["Analyst"])
        result = _derive_roles(draft)
        assert result.count("Analyst") == 1



# ---------------------------------------------------------------------------

class TestBuildExportDocx:
    def _bundle(self):
        return build_evidence_bundle(DRAFT_WITH_ANCHORS, JOB_WITH_SIGNALS)

    def test_returns_bytes(self):
        result = build_export_docx(DRAFT_WITH_ANCHORS, self._bundle(), "job-123")
        assert isinstance(result, bytes)

    def test_docx_magic_bytes(self):
        """DOCX files are ZIP archives starting with PK."""
        result = build_export_docx(DRAFT_WITH_ANCHORS, self._bundle(), "job-123")
        assert result[:2] == b"PK"

    def test_valid_docx_readable(self):
        """Resulting bytes should be parseable by python-docx."""
        import io
        from docx import Document

        result = build_export_docx(DRAFT_WITH_ANCHORS, self._bundle(), "job-123")
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Process Definition Document" in full_text or any(
            "Process Definition Document" in para.text for para in doc.paragraphs
        )

    def test_docx_exceptions_table_uses_trigger_field(self):
        import io
        from docx import Document

        draft = {
            "pdd": {
                "purpose": "p",
                "scope": "s",
                "steps": [],
                "exceptions": [
                    {
                        "scenario": "Validation failed",
                        "trigger": "Missing mandatory field",
                        "action_required": "Request correction",
                        "owner": "Analyst",
                    }
                ],
            },
            "sipoc": [],
        }
        result = build_export_docx(draft, build_evidence_bundle(draft, JOB_MINIMAL), "job-1")
        doc = Document(io.BytesIO(result))
        exception_table = None
        for table in doc.tables:
            if table.rows and table.rows[0].cells and table.rows[0].cells[0].text == "Exception Scenario":
                exception_table = table
                break
        assert exception_table is not None
        assert exception_table.rows[1].cells[1].text == "Missing mandatory field"

    def test_docx_raci_uses_approval_matrix_for_non_actor_roles(self):
        import io
        from docx import Document

        draft = {
            "pdd": {
                "purpose": "p",
                "scope": "s",
                "steps": [
                    {
                        "id": "step-01",
                        "summary": "Capture request",
                        "actor": "Analyst",
                        "source_anchors": [{"anchor": "00:00:01-00:00:05", "confidence": 0.9}],
                    }
                ],
                "roles": ["Analyst", "Manager", "Compliance"],
                "exceptions": [],
            },
            "approval_matrix": [
                {"role": "Manager", "responsibility": "A"},
                {"role": "Compliance", "responsibility": "C"},
            ],
            "sipoc": [],
        }
        result = build_export_docx(draft, build_evidence_bundle(draft, JOB_MINIMAL), "job-raci")
        doc = Document(io.BytesIO(result))
        raci_table = None
        for table in doc.tables:
            if table.rows and table.rows[0].cells and table.rows[0].cells[0].text == "Task":
                raci_table = table
                break
        assert raci_table is not None
        assert [cell.text for cell in raci_table.rows[1].cells] == [
            "Capture request",
            "R",
            "A",
            "C",
        ]

    def test_contains_sipoc_table(self):
        import io
        from docx import Document

        result = build_export_docx(DRAFT_WITH_ANCHORS, self._bundle(), "job-123")
        doc = Document(io.BytesIO(result))
        # At least one table (SIPOC) should be present
        assert len(doc.tables) >= 1

    def test_contains_evidence_bundle_table(self):
        import io
        from docx import Document

        result = build_export_docx(DRAFT_WITH_ANCHORS, self._bundle(), "job-123")
        doc = Document(io.BytesIO(result))
        # SIPOC table + evidence bundle table = at least 2 tables
        assert len(doc.tables) >= 2

    def test_empty_draft_does_not_crash(self):
        bundle = build_evidence_bundle({}, {})
        result = build_export_docx({}, bundle, "job-empty")
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_no_sipoc_fallback_paragraph(self):
        import io
        from docx import Document

        draft_no_sipoc = {
            "pdd": {"purpose": "p", "scope": "s", "steps": []},
            "sipoc": [],
        }
        bundle = build_evidence_bundle(draft_no_sipoc, JOB_MINIMAL)
        result = build_export_docx(draft_no_sipoc, bundle, "job-x")
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "No SIPOC rows available" in full_text

    def test_docx_embeds_frame_from_bytes_map(self):
        bundle = {
            "linked_anchors": [],
            "evidence_strength": "high",
            "frame_captures_note": "",
            "frame_captures": [{"storage_key": "job-123/frames/frame_0001.jpg", "timestamp_sec": 12.0}],
        }
        result = build_export_docx(
            DRAFT_WITH_ANCHORS,
            bundle,
            "job-123",
            frame_bytes_map={"job-123/frames/frame_0001.jpg": TINY_PNG_BYTES},
        )
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_docx_skips_missing_frame_gracefully(self):
        bundle = {
            "linked_anchors": [],
            "evidence_strength": "high",
            "frame_captures_note": "",
            "frame_captures": [{"storage_key": "job-123/frames/missing.jpg", "timestamp_sec": 12.0}],
        }
        result = build_export_docx(DRAFT_WITH_ANCHORS, bundle, "job-123", frame_bytes_map={})
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_docx_semantic_review_section_hidden_when_empty(self):
        import io
        from docx import Document

        result = build_export_docx(DRAFT_WITH_ANCHORS, self._bundle(), "job-123")
        doc = Document(io.BytesIO(result))
        assert not any(p.text.strip() == "5A. AI Semantic Review Notes" for p in doc.paragraphs)

    def test_docx_semantic_review_section_rendered_when_flags_present(self):
        import io
        from docx import Document

        bundle = self._bundle()
        bundle["llm_semantic_flags"] = [
            {
                "code": "factual_drift",
                "message": "Actor mismatch against cited evidence.",
                "evidence_id": "ev-02",
                "step_id": "step-02",
                "anchor": "00:01:15",
                "snippet": "Manager receives a copy only",
            }
        ]
        result = build_export_docx(DRAFT_WITH_ANCHORS, bundle, "job-123")
        doc = Document(io.BytesIO(result))
        assert any(p.text.strip() == "5A. AI Semantic Review Notes" for p in doc.paragraphs)
        assert any("ev-02" in cell.text for table in doc.tables for row in table.rows for cell in row.cells)
