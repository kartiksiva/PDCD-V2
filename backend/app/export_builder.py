"""Evidence-linked export builders for PDF, DOCX, and Markdown (PRD §8.10)."""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from app.agents.anchor_utils import classify_anchor


@dataclass
class _AnchorEntry:
    anchor_id: str
    anchor_type: str  # timestamp_range | frame_id | section_label | missing
    anchor_value: str
    confidence: float
    linked_step_ids: List[str] = field(default_factory=list)
    linked_sipoc_rows: List[int] = field(default_factory=list)
    ocr_snippet: Optional[str] = None


def _classify_anchor_type(anchor: str) -> str:
    """Classify anchor string into a type label."""
    return classify_anchor(anchor)


def _timestamp_to_seconds(value: str) -> float:
    parts = value.split(":")
    try:
        total = 0.0
        for index, part in enumerate(parts):
            total += float(part) * (60 ** (len(parts) - 1 - index))
        return total
    except (TypeError, ValueError):
        return -1.0


def _is_safe_storage_key(storage_key: str) -> bool:
    normalized = (storage_key or "").strip().replace("\\", "/")
    if not normalized:
        return False
    if normalized.startswith("/") or normalized.startswith("./"):
        return False
    parts = normalized.split("/")
    if any(part in ("", ".", "..") for part in parts):
        return False
    return True


def build_evidence_bundle(finalized_draft: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    """Build evidence bundle manifest from finalized draft and job signals.

    Per PRD §8.10:
    - Only anchors linked to at least one PDD step or SIPOC row are included.
    - OCR snippets attached when available from extracted_evidence.
    - Frame captures are noted as pending (VideoAdapter stub).
    """
    pdd_steps: List[Dict[str, Any]] = (finalized_draft.get("pdd") or {}).get("steps") or []
    sipoc_rows: List[Dict[str, Any]] = finalized_draft.get("sipoc") or []

    anchor_map: Dict[str, _AnchorEntry] = {}

    # Collect anchors from PDD steps (each step may have multiple source_anchors)
    for step in pdd_steps:
        step_id = step.get("id", "")
        for sa in step.get("source_anchors") or []:
            anchor_val = sa.get("anchor", "")
            if not anchor_val:
                continue
            if anchor_val not in anchor_map:
                anchor_map[anchor_val] = _AnchorEntry(
                    anchor_id=f"a-{len(anchor_map) + 1}",
                    anchor_type=_classify_anchor_type(anchor_val),
                    anchor_value=anchor_val,
                    confidence=float(sa.get("confidence") or 0.0),
                )
            entry = anchor_map[anchor_val]
            if step_id and step_id not in entry.linked_step_ids:
                entry.linked_step_ids.append(step_id)

    # Collect anchors from SIPOC rows
    for idx, row in enumerate(sipoc_rows):
        anchor_val = row.get("source_anchor", "")
        if not anchor_val:
            continue
        if anchor_val not in anchor_map:
            anchor_map[anchor_val] = _AnchorEntry(
                anchor_id=f"a-{len(anchor_map) + 1}",
                anchor_type=_classify_anchor_type(anchor_val),
                anchor_value=anchor_val,
                confidence=0.0,
            )
        entry = anchor_map[anchor_val]
        if idx not in entry.linked_sipoc_rows:
            entry.linked_sipoc_rows.append(idx)

    # Attach OCR snippets from extracted evidence when anchor values match
    evidence_items: List[Dict[str, Any]] = (
        (job.get("extracted_evidence") or {}).get("evidence_items") or []
    )
    for item in evidence_items:
        # Extraction schema uses "anchor"; "source_anchor" is kept as fallback.
        item_anchor = item.get("anchor") or item.get("source_anchor") or ""
        if item_anchor and item_anchor in anchor_map:
            anchor_map[item_anchor].ocr_snippet = (
                item.get("ocr_text") or item.get("content_snippet")
            )

    # Collect frame captures from evidence metadata or persisted agent signals.
    frame_captures: list[dict] = []
    for item in evidence_items:
        frame_keys = (item.get("metadata") or {}).get("frame_storage_keys") or []
        for storage_key, timestamp_sec in frame_keys:
            if not _is_safe_storage_key(storage_key):
                continue
            frame_captures.append({"storage_key": storage_key, "timestamp_sec": timestamp_sec})
    for storage_key, timestamp_sec in (job.get("agent_signals") or {}).get("frame_storage_keys") or []:
        if not _is_safe_storage_key(storage_key):
            continue
        if not any(
            existing["storage_key"] == storage_key and existing["timestamp_sec"] == timestamp_sec
            for existing in frame_captures
        ):
            frame_captures.append({"storage_key": storage_key, "timestamp_sec": timestamp_sec})

    # Link frame captures to timestamp anchors by midpoint proximity.
    for capture in frame_captures:
        timestamp_sec = capture["timestamp_sec"]
        for anchor_val, entry in anchor_map.items():
            if entry.anchor_type != "timestamp_range":
                continue
            parts = anchor_val.split("-")
            if len(parts) == 1:
                midpoint = _timestamp_to_seconds(parts[0])
            else:
                start = _timestamp_to_seconds(parts[0])
                end = _timestamp_to_seconds(parts[-1])
                if start < 0 or end < 0:
                    continue
                midpoint = (start + end) / 2
            if midpoint >= 0 and abs(timestamp_sec - midpoint) <= 10:
                capture.setdefault("linked_anchor_ids", []).append(entry.anchor_id)

    # PRD §8.10: only include anchors linked to at least one step or SIPOC row
    linked_anchors = [
        {
            "anchor_id": e.anchor_id,
            "anchor_type": e.anchor_type,
            "anchor_value": e.anchor_value,
            "confidence": e.confidence,
            "linked_step_ids": e.linked_step_ids,
            "linked_sipoc_rows": e.linked_sipoc_rows,
            "ocr_snippet": e.ocr_snippet,
        }
        for e in anchor_map.values()
        if e.linked_step_ids or e.linked_sipoc_rows
    ]

    return {
        "evidence_strength": (job.get("agent_signals") or {}).get("evidence_strength"),
        "frame_policy": (
            (job.get("input_manifest") or {}).get("video", {}).get("frame_extraction_policy")
        ),
        "linked_anchors": linked_anchors,
        "total_linked_anchors": len(linked_anchors),
        "frame_captures": frame_captures,
        "frame_captures_note": (
            "Frame captures embedded above."
            if frame_captures
            else "No frame captures available for this job."
        ),
        "unlinked_note": "Evidence not included in this export format",
    }


def _needs_review(value: Any) -> str:
    if value is None:
        return "Needs Review"
    if isinstance(value, str):
        text = value.strip()
        return text or "Needs Review"
    return str(value)


def _format_date(value: Any) -> str:
    """Format an ISO timestamp or datetime to DD-MMM-YYYY for SOP headers."""
    if value is None:
        return "Needs Review"
    from datetime import datetime
    s = str(value).strip()
    for fmt in ("%Y-%m-%dT%H:%M:%S.%f%z", "%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%dT%H:%M:%SZ",
                "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(s[:26].rstrip("Z"), fmt.rstrip("z").rstrip("%z")).strftime("%d-%b-%Y")
        except ValueError:
            continue
    return s or "Needs Review"


def _derive_roles(draft: Dict[str, Any]) -> List[str]:
    pdd = draft.get("pdd") or {}
    roles: List[str] = []
    for role in pdd.get("roles") or []:
        role_name = str(role).strip()
        if role_name and role_name not in roles:
            roles.append(role_name)
    for step in pdd.get("steps") or []:
        actor = str((step or {}).get("actor") or "").strip()
        if actor and actor not in roles:
            roles.append(actor)
    return roles or ["Needs Review"]


def _step_source_anchor(step: Dict[str, Any]) -> str:
    anchors = [
        (sa or {}).get("anchor", "").strip()
        for sa in (step.get("source_anchors") or [])
        if (sa or {}).get("anchor")
    ]
    return ", ".join(anchors) if anchors else "Needs Review"


def _step_tools_systems(step: Dict[str, Any]) -> str:
    return _needs_review(step.get("tools_systems") or step.get("system"))


def build_export_markdown(draft: Dict[str, Any], evidence_bundle: Dict[str, Any]) -> str:
    """Build SOP-style Markdown export."""
    if not draft:
        return "No finalized draft available."

    pdd = draft.get("pdd") or {}
    steps = pdd.get("steps") or []
    sipoc_rows = draft.get("sipoc") or []
    roles = _derive_roles(draft)
    approval_rows = draft.get("approval_matrix") or []
    controls = pdd.get("process_controls") or draft.get("process_controls") or []
    exceptions = pdd.get("exceptions") or []
    automation = draft.get("automation_opportunities") or []
    faqs = draft.get("faqs") or []
    process_name = _needs_review(
        pdd.get("process_name") or draft.get("subject_process") or pdd.get("purpose")
    )

    parts = [
        "# Standard Operating Procedure (SOP)",
        "",
        f"## {process_name}",
        f"**Function:** {_needs_review(pdd.get('function'))}",
        f"**Sub-Function:** {_needs_review(pdd.get('sub_function'))}",
        "**Document Version:** v1.0",
        "**Document Status:** Draft",
        f"**Effective Date:** {_format_date(draft.get('generated_at'))}",
        "",
        "---",
        "",
        "## 1. Document Control",
        "### 1.1 Key Stakeholders",
        "| # | Name | Position / Designation | Email ID |",
        "|---|------|------------------------|----------|",
        "| 1 | Needs Review | Needs Review | Needs Review |",
        "",
        "### 1.2 Version History",
        "| Version | Date | Status | Author | Reviewed By | Comments / Changes |",
        "|---------|------|--------|--------|-------------|-------------------|",
        f"| v1.0 | {_format_date(draft.get('generated_at'))} | Draft | Needs Review | Needs Review | Initial export |",
        "",
        "---",
        "",
        "## Index",
        "1. Document Control  2. Introduction  3. Process Steps  4. Process Exceptions  5. Process Controls  6. Approval Matrix  7. Appendix",
        "",
        "---",
        "",
        "## 2. Introduction",
        "### 2.1 Process Overview",
        _needs_review(pdd.get("process_overview") or pdd.get("scope")),
        "",
        "### 2.2 Process Objective",
        _needs_review(pdd.get("process_objective") or pdd.get("purpose")),
        "",
        "### 2.3 Frequency",
        _needs_review(pdd.get("frequency")),
        "",
        "### 2.4 SLA",
        _needs_review(pdd.get("sla")),
        "",
        "### 2.5 RACI (task × role matrix)",
    ]

    raci_header = "| Task | " + " | ".join(roles) + " |"
    raci_sep = "|------|" + "|".join("---" for _ in roles) + "|"
    parts.extend([raci_header, raci_sep])
    for step in steps:
        actor = str((step or {}).get("actor") or "").strip()
        task = _needs_review((step or {}).get("summary") or (step or {}).get("id"))
        row = [task]
        for role in roles:
            if actor and role == actor:
                row.append("R")
            elif actor:
                row.append("—")
            else:
                row.append("Needs Review")
        parts.append("| " + " | ".join(row) + " |")

    if not steps:
        parts.append("| Needs Review | " + " | ".join("Needs Review" for _ in roles) + " |")

    parts.extend(["", "### 2.6 SIPOC (Supplier / Input / Process / Output / Customer)"])
    parts.extend(
        [
            "| Supplier | Input | Process | Output | Customer | Step Anchor | Source Anchor |",
            "|----------|-------|---------|--------|----------|-------------|---------------|",
        ]
    )
    if sipoc_rows:
        for row in sipoc_rows:
            step_anchor = ", ".join((row.get("step_anchor") or [])) or "Needs Review"
            parts.append(
                "| "
                + " | ".join(
                    [
                        _needs_review(row.get("supplier")),
                        _needs_review(row.get("input")),
                        _needs_review(row.get("process_step")),
                        _needs_review(row.get("output")),
                        _needs_review(row.get("customer")),
                        step_anchor,
                        _needs_review(row.get("source_anchor")),
                    ]
                )
                + " |"
            )
    else:
        parts.append("| Needs Review | Needs Review | Needs Review | Needs Review | Needs Review | Needs Review | Needs Review |")

    parts.extend(["", "---", "", "## 3. Process Steps"])
    if steps:
        for idx, step in enumerate(steps, start=1):
            parts.extend(
                [
                    f"### Step {idx}: {_needs_review(step.get('summary') or step.get('id'))} ({_needs_review(step.get('id'))})",
                    f"- Description: {_needs_review(step.get('summary'))}",
                    f"- Tools / Systems: {_step_tools_systems(step)}",
                    f"- Inputs / Outputs: {_needs_review(step.get('input'))} -> {_needs_review(step.get('output'))}",
                    f"- Source Timestamp (evidence anchor): {_step_source_anchor(step)}",
                    "",
                ]
            )
    else:
        parts.extend(
            [
                "### Step 1: Needs Review",
                "- Description: Needs Review",
                "- Tools / Systems: Needs Review",
                "- Inputs / Outputs: Needs Review -> Needs Review",
                "- Source Timestamp (evidence anchor): Needs Review",
                "",
            ]
        )

    parts.extend(["---", "", "## 4. Process Exceptions"])
    parts.extend(
        [
            "| Exception Scenario | Description | Action Required | Owner |",
            "|--------------------|-------------|-----------------|-------|",
        ]
    )
    if exceptions:
        for exc in exceptions:
            if isinstance(exc, dict):
                parts.append(
                    "| "
                    + " | ".join(
                        [
                            _needs_review(exc.get("scenario")),
                            _needs_review(exc.get("description")),
                            _needs_review(exc.get("action_required")),
                            _needs_review(exc.get("owner")),
                        ]
                    )
                    + " |"
                )
            else:
                parts.append(f"| {_needs_review(exc)} | {_needs_review(exc)} | Needs Review | Needs Review |")
    else:
        parts.append("| Needs Review | Needs Review | Needs Review | Needs Review |")

    parts.extend(["", "## 5. Process Controls"])
    parts.extend(
        [
            "| Control # | Process Step | Control Description | Manual/System | Preventive/Detective |",
            "|-----------|--------------|---------------------|---------------|----------------------|",
        ]
    )
    if controls:
        for idx, ctl in enumerate(controls, start=1):
            if isinstance(ctl, dict):
                parts.append(
                    "| "
                    + " | ".join(
                        [
                            _needs_review(ctl.get("control_id") or f"control-{idx:02d}"),
                            _needs_review(ctl.get("process_step_id") or ctl.get("process_step")),
                            _needs_review(ctl.get("control_description")),
                            _needs_review(ctl.get("manual_or_system")),
                            _needs_review(ctl.get("preventive_or_detective")),
                        ]
                    )
                    + " |"
                )
            else:
                parts.append(f"| control-{idx:02d} | Needs Review | {_needs_review(ctl)} | Needs Review | Needs Review |")
    else:
        parts.append("| control-01 | Needs Review | Needs Review | Needs Review | Needs Review |")

    parts.extend(["", "## 6. Approval Matrix"])
    parts.extend(["| Role | Responsibility |", "|------|----------------|"])
    if approval_rows:
        for row in approval_rows:
            if isinstance(row, dict):
                parts.append(
                    f"| {_needs_review(row.get('role'))} | {_needs_review(row.get('responsibility'))} |"
                )
            else:
                parts.append(f"| {_needs_review(row)} | Needs Review |")
    else:
        for role in roles:
            parts.append(f"| {role} | Needs Review |")

    parts.extend(["", "## 7. Appendix", "### Automation Opportunities"])
    parts.extend(
        [
            "| ID | Description | Quantification | Automation Signal |",
            "|----|-------------|----------------|-------------------|",
        ]
    )
    if automation:
        for idx, item in enumerate(automation, start=1):
            if isinstance(item, dict):
                parts.append(
                    "| "
                    + " | ".join(
                        [
                            _needs_review(item.get("id") or f"auto-{idx:02d}"),
                            _needs_review(item.get("description")),
                            _needs_review(item.get("quantification")),
                            _needs_review(item.get("automation_signal")),
                        ]
                    )
                    + " |"
                )
            else:
                parts.append(f"| auto-{idx:02d} | {_needs_review(item)} | Needs Review | Needs Review |")
    else:
        parts.append("| auto-01 | Needs Review | Needs Review | Needs Review |")

    parts.extend(["", "### FAQs"])
    if faqs:
        for idx, faq in enumerate(faqs, start=1):
            if isinstance(faq, dict):
                parts.append(f"{idx}. **Q:** {_needs_review(faq.get('question'))}")
                parts.append(f"   **A:** {_needs_review(faq.get('answer'))}")
            else:
                parts.append(f"{idx}. {_needs_review(faq)}")
    else:
        parts.append("1. Needs Review")

    # Retain evidence manifest for anchor traceability.
    parts.extend(["", "### Evidence Bundle Manifest", ""])
    parts.append(f"**Evidence Strength:** {evidence_bundle.get('evidence_strength') or 'unknown'}")
    linked_anchors = evidence_bundle.get("linked_anchors") or []
    if linked_anchors:
        parts.extend(
            [
                "",
                "| Anchor | Type | Confidence | Linked Steps | OCR Snippet |",
                "|--------|------|------------|-------------|-------------|",
            ]
        )
        for a in linked_anchors:
            steps_str = ", ".join(a.get("linked_step_ids") or []) or "—"
            ocr = (a.get("ocr_snippet") or "—")[:80]
            conf = f"{a.get('confidence', 0.0):.2f}"
            parts.append(
                f"| `{a['anchor_value']}` | {a['anchor_type']} | {conf} | {steps_str} | {ocr} |"
            )
    else:
        parts.append("*No linked evidence anchors found.*")

    note = evidence_bundle.get("frame_captures_note", "")
    if note:
        parts.extend(["", f"> {note}"])
    frame_captures = evidence_bundle.get("frame_captures") or []
    parts.extend(["", "### Frame captures"])
    if frame_captures:
        for capture in frame_captures:
            anchor_ids = ", ".join(capture.get("linked_anchor_ids") or []) or "unlinked"
            parts.append(
                f"- `{capture.get('storage_key')}` @ {capture.get('timestamp_sec', 0):.1f}s "
                f"(anchors: {anchor_ids})"
            )
    else:
        parts.append("- No frame captures available.")

    return "\n".join(parts)


def build_export_pdf(
    draft: Dict[str, Any],
    evidence_bundle: Dict[str, Any],
    frame_bytes_map: Optional[Dict[str, bytes]] = None,
) -> bytes:
    """Build evidence-linked PDF export."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.add_page()
    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    def _heading(text: str, size: int = 12, bold: bool = False) -> None:
        pdf.set_x(pdf.l_margin)
        style = "B" if bold else ""
        pdf.set_font("Helvetica", style=style, size=size)
        pdf.cell(page_width, 10, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", size=11)

    def _row(text: str) -> None:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_width, 7, text)

    pdf.set_font("Helvetica", size=11)
    pdd = draft.get("pdd") or {}

    _heading("Process Definition Document", size=14, bold=True)
    pdf.ln(2)
    _row(f"Purpose: {pdd.get('purpose', '')}")
    _row(f"Scope: {pdd.get('scope', '')}")

    pdf.ln(3)
    _heading("Steps", bold=True)
    for step in pdd.get("steps") or []:
        anchors = ", ".join(
            sa.get("anchor", "")
            for sa in (step.get("source_anchors") or [])
            if sa.get("anchor")
        )
        anchor_part = f" [anchors: {anchors}]" if anchors else ""
        _row(f"  {step.get('id')}: {step.get('summary')}{anchor_part}")

    pdf.ln(3)
    _heading("SIPOC", bold=True)
    for idx, row in enumerate(draft.get("sipoc") or [], start=1):
        step_refs = ", ".join(row.get("step_anchor") or [])
        step_ref_part = f" [steps: {step_refs}]" if step_refs else ""
        _row(
            f"  {idx}. {row.get('process_step')} "
            f"[anchor: {row.get('source_anchor', 'N/A')}]{step_ref_part}"
        )

    # Evidence bundle section
    pdf.ln(4)
    _heading("Evidence Bundle", bold=True)
    strength = evidence_bundle.get("evidence_strength") or "unknown"
    _row(f"Evidence Strength: {strength}")

    linked_anchors = evidence_bundle.get("linked_anchors") or []
    if linked_anchors:
        pdf.ln(2)
        _row(f"Total linked anchors: {len(linked_anchors)}")
        pdf.ln(1)
        for a in linked_anchors:
            steps_str = ", ".join(a.get("linked_step_ids") or []) or "none"
            conf = f"{a.get('confidence', 0.0):.2f}"
            _row(
                f"  [{a['anchor_type']}] {a['anchor_value']}  "
                f"confidence: {conf}  steps: {steps_str}"
            )
            if a.get("ocr_snippet"):
                snippet = a["ocr_snippet"][:120]
                _row(f"    OCR: {snippet}")
    else:
        _row("No linked evidence anchors found.")

    note = evidence_bundle.get("frame_captures_note", "")
    if note:
        pdf.ln(2)
        _row(f"Note: {note}")

    frame_captures = evidence_bundle.get("frame_captures") or []
    if frame_captures:
        pdf.ln(3)
        _heading("Frame Captures", bold=True)
        for capture in frame_captures:
            key = capture.get("storage_key", "")
            timestamp_sec = capture.get("timestamp_sec", 0.0)
            image_bytes = (frame_bytes_map or {}).get(key)
            if image_bytes:
                try:
                    pdf.image(io.BytesIO(image_bytes), w=120)
                    pdf.set_font("Helvetica", size=8)
                    pdf.cell(0, 5, f"Frame @ {timestamp_sec:.1f}s — {key}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.set_font("Helvetica", size=11)
                except Exception:
                    _row(f"Frame @ {timestamp_sec:.1f}s (image unreadable): {key}")
            else:
                pdf.set_font("Helvetica", size=8)
                pdf.cell(
                    0,
                    5,
                    f"Frame @ {timestamp_sec:.1f}s: {key} (not available)",
                    new_x=XPos.LMARGIN,
                    new_y=YPos.NEXT,
                )
                pdf.set_font("Helvetica", size=11)

    return bytes(pdf.output())


def build_export_docx(
    draft: Dict[str, Any],
    evidence_bundle: Dict[str, Any],
    job_id: str,
    frame_bytes_map: Optional[Dict[str, bytes]] = None,
) -> bytes:
    """Build SOP-style DOCX export using python-docx."""
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    pdd = draft.get("pdd") or {}
    steps = pdd.get("steps") or []
    roles = _derive_roles(draft)
    sipoc = draft.get("sipoc") or []
    controls = pdd.get("process_controls") or draft.get("process_controls") or []
    exceptions = pdd.get("exceptions") or []
    approval_rows = draft.get("approval_matrix") or []
    automation = draft.get("automation_opportunities") or []
    faqs = draft.get("faqs") or []

    process_name = _needs_review(
        pdd.get("process_name") or draft.get("subject_process") or pdd.get("purpose")
    )

    doc.add_heading("Standard Operating Procedure (SOP)", level=0)
    doc.add_heading(process_name, level=1)
    doc.add_paragraph(f"Function: {_needs_review(pdd.get('function'))}")
    doc.add_paragraph(f"Sub-Function: {_needs_review(pdd.get('sub_function'))}")
    doc.add_paragraph("Document Version: v1.0")
    doc.add_paragraph("Document Status: Draft")
    doc.add_paragraph(f"Effective Date: {_format_date(draft.get('generated_at'))}")

    doc.add_heading("1. Document Control", level=1)
    doc.add_heading("1.1 Key Stakeholders", level=2)
    stakeholders = doc.add_table(rows=1, cols=4)
    stakeholders.style = "Table Grid"
    for i, h in enumerate(["#", "Name", "Position / Designation", "Email ID"]):
        stakeholders.rows[0].cells[i].text = h
    row = stakeholders.add_row().cells
    row[0].text = "1"
    row[1].text = "Needs Review"
    row[2].text = "Needs Review"
    row[3].text = "Needs Review"

    doc.add_heading("1.2 Version History", level=2)
    vh = doc.add_table(rows=1, cols=6)
    vh.style = "Table Grid"
    for i, h in enumerate(["Version", "Date", "Status", "Author", "Reviewed By", "Comments / Changes"]):
        vh.rows[0].cells[i].text = h
    row = vh.add_row().cells
    row[0].text = "v1.0"
    row[1].text = _format_date(draft.get("generated_at"))
    row[2].text = "Draft"
    row[3].text = "Needs Review"
    row[4].text = "Needs Review"
    row[5].text = "Initial export"

    doc.add_heading("Index", level=1)
    doc.add_paragraph("1. Document Control  2. Introduction  3. Process Steps  4. Process Exceptions  5. Process Controls  6. Approval Matrix  7. Appendix")

    doc.add_heading("2. Introduction", level=1)
    doc.add_heading("2.1 Process Overview", level=2)
    doc.add_paragraph(_needs_review(pdd.get("process_overview") or pdd.get("scope")))
    doc.add_heading("2.2 Process Objective", level=2)
    doc.add_paragraph(_needs_review(pdd.get("process_objective") or pdd.get("purpose")))
    doc.add_heading("2.3 Frequency", level=2)
    doc.add_paragraph(_needs_review(pdd.get("frequency")))
    doc.add_heading("2.4 SLA", level=2)
    doc.add_paragraph(_needs_review(pdd.get("sla")))

    doc.add_heading("2.5 RACI (task × role matrix)", level=2)
    raci = doc.add_table(rows=1, cols=1 + len(roles))
    raci.style = "Table Grid"
    raci.rows[0].cells[0].text = "Task"
    for i, role in enumerate(roles, start=1):
        raci.rows[0].cells[i].text = role
    if steps:
        for step in steps:
            actor = str((step or {}).get("actor") or "").strip()
            row = raci.add_row().cells
            row[0].text = _needs_review((step or {}).get("summary") or (step or {}).get("id"))
            for i, role in enumerate(roles, start=1):
                if actor and role == actor:
                    row[i].text = "R"
                elif actor:
                    row[i].text = "—"
                else:
                    row[i].text = "Needs Review"
    else:
        row = raci.add_row().cells
        row[0].text = "Needs Review"
        for i in range(1, len(roles) + 1):
            row[i].text = "Needs Review"

    doc.add_heading("2.6 SIPOC (Supplier / Input / Process / Output / Customer)", level=2)
    if sipoc:
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(["Supplier", "Input", "Process", "Output", "Customer", "Step Anchor", "Source Anchor"]):
            hdr_cells[i].text = h
        for row_data in sipoc:
            cells = table.add_row().cells
            cells[0].text = _needs_review(row_data.get("supplier"))
            cells[1].text = _needs_review(row_data.get("input"))
            cells[2].text = _needs_review(row_data.get("process_step"))
            cells[3].text = _needs_review(row_data.get("output"))
            cells[4].text = _needs_review(row_data.get("customer"))
            cells[5].text = ", ".join(row_data.get("step_anchor") or []) or "Needs Review"
            cells[6].text = _needs_review(row_data.get("source_anchor"))
    else:
        doc.add_paragraph("No SIPOC rows available.")

    doc.add_heading("3. Process Steps", level=1)
    if steps:
        for idx, step in enumerate(steps, start=1):
            doc.add_heading(
                f"Step {idx}: {_needs_review(step.get('summary') or step.get('id'))}",
                level=2,
            )
            doc.add_paragraph(f"Description: {_needs_review(step.get('summary'))}")
            doc.add_paragraph(f"Tools / Systems: {_step_tools_systems(step)}")
            doc.add_paragraph(
                "Inputs / Outputs: "
                f"{_needs_review(step.get('input'))} -> {_needs_review(step.get('output'))}"
            )
            doc.add_paragraph(
                "Source Timestamp (evidence anchor): "
                f"{_step_source_anchor(step)}"
            )
    else:
        doc.add_paragraph("Step 1: Needs Review")
        doc.add_paragraph("Description: Needs Review")
        doc.add_paragraph("Tools / Systems: Needs Review")
        doc.add_paragraph("Inputs / Outputs: Needs Review -> Needs Review")
        doc.add_paragraph("Source Timestamp (evidence anchor): Needs Review")

    doc.add_heading("4. Process Exceptions", level=1)
    ex_table = doc.add_table(rows=1, cols=4)
    ex_table.style = "Table Grid"
    for i, h in enumerate(["Exception Scenario", "Description", "Action Required", "Owner"]):
        ex_table.rows[0].cells[i].text = h
    if exceptions:
        for exc in exceptions:
            row = ex_table.add_row().cells
            if isinstance(exc, dict):
                row[0].text = _needs_review(exc.get("scenario"))
                row[1].text = _needs_review(exc.get("description"))
                row[2].text = _needs_review(exc.get("action_required"))
                row[3].text = _needs_review(exc.get("owner"))
            else:
                row[0].text = _needs_review(exc)
                row[1].text = _needs_review(exc)
                row[2].text = "Needs Review"
                row[3].text = "Needs Review"
    else:
        row = ex_table.add_row().cells
        row[0].text = "Needs Review"
        row[1].text = "Needs Review"
        row[2].text = "Needs Review"
        row[3].text = "Needs Review"

    doc.add_heading("5. Process Controls", level=1)
    ctl_table = doc.add_table(rows=1, cols=5)
    ctl_table.style = "Table Grid"
    for i, h in enumerate(["Control #", "Process Step", "Control Description", "Manual/System", "Preventive/Detective"]):
        ctl_table.rows[0].cells[i].text = h
    if controls:
        for idx, ctl in enumerate(controls, start=1):
            row = ctl_table.add_row().cells
            if isinstance(ctl, dict):
                row[0].text = _needs_review(ctl.get("control_id") or f"control-{idx:02d}")
                row[1].text = _needs_review(ctl.get("process_step_id") or ctl.get("process_step"))
                row[2].text = _needs_review(ctl.get("control_description"))
                row[3].text = _needs_review(ctl.get("manual_or_system"))
                row[4].text = _needs_review(ctl.get("preventive_or_detective"))
            else:
                row[0].text = f"control-{idx:02d}"
                row[1].text = "Needs Review"
                row[2].text = _needs_review(ctl)
                row[3].text = "Needs Review"
                row[4].text = "Needs Review"
    else:
        row = ctl_table.add_row().cells
        row[0].text = "control-01"
        row[1].text = "Needs Review"
        row[2].text = "Needs Review"
        row[3].text = "Needs Review"
        row[4].text = "Needs Review"

    doc.add_heading("6. Approval Matrix", level=1)
    appr_table = doc.add_table(rows=1, cols=2)
    appr_table.style = "Table Grid"
    appr_table.rows[0].cells[0].text = "Role"
    appr_table.rows[0].cells[1].text = "Responsibility"
    if approval_rows:
        for row_data in approval_rows:
            row = appr_table.add_row().cells
            if isinstance(row_data, dict):
                row[0].text = _needs_review(row_data.get("role"))
                row[1].text = _needs_review(row_data.get("responsibility"))
            else:
                row[0].text = _needs_review(row_data)
                row[1].text = "Needs Review"
    else:
        for role in roles:
            row = appr_table.add_row().cells
            row[0].text = role
            row[1].text = "Needs Review"

    doc.add_heading("7. Appendix", level=1)
    doc.add_heading("Automation Opportunities", level=2)
    auto_table = doc.add_table(rows=1, cols=4)
    auto_table.style = "Table Grid"
    for i, h in enumerate(["ID", "Description", "Quantification", "Automation Signal"]):
        auto_table.rows[0].cells[i].text = h
    if automation:
        for idx, item in enumerate(automation, start=1):
            row = auto_table.add_row().cells
            if isinstance(item, dict):
                row[0].text = _needs_review(item.get("id") or f"auto-{idx:02d}")
                row[1].text = _needs_review(item.get("description"))
                row[2].text = _needs_review(item.get("quantification"))
                row[3].text = _needs_review(item.get("automation_signal"))
            else:
                row[0].text = f"auto-{idx:02d}"
                row[1].text = _needs_review(item)
                row[2].text = "Needs Review"
                row[3].text = "Needs Review"
    else:
        row = auto_table.add_row().cells
        row[0].text = "auto-01"
        row[1].text = "Needs Review"
        row[2].text = "Needs Review"
        row[3].text = "Needs Review"

    doc.add_heading("FAQs", level=2)
    if faqs:
        for idx, faq in enumerate(faqs, start=1):
            if isinstance(faq, dict):
                doc.add_paragraph(f"{idx}. Q: {_needs_review(faq.get('question'))}")
                doc.add_paragraph(f"   A: {_needs_review(faq.get('answer'))}")
            else:
                doc.add_paragraph(f"{idx}. {_needs_review(faq)}")
    else:
        doc.add_paragraph("1. Needs Review")

    # Keep evidence manifest in the appendix for traceability.
    doc.add_heading("Evidence Bundle Manifest", level=2)
    strength = evidence_bundle.get("evidence_strength") or "unknown"
    doc.add_paragraph(f"Evidence Strength: {strength}")
    linked_anchors = evidence_bundle.get("linked_anchors") or []
    if linked_anchors:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for i, h in enumerate(["Anchor", "Type", "Confidence", "Linked Steps", "OCR Snippet"]):
            table.rows[0].cells[i].text = h
        for a in linked_anchors:
            row = table.add_row().cells
            row[0].text = a["anchor_value"]
            row[1].text = a["anchor_type"]
            row[2].text = f"{a.get('confidence', 0.0):.2f}"
            row[3].text = ", ".join(a.get("linked_step_ids") or []) or "—"
            row[4].text = (a.get("ocr_snippet") or "—")[:120]
    else:
        doc.add_paragraph("No linked evidence anchors found.")

    frame_captures = evidence_bundle.get("frame_captures") or []
    if frame_captures:
        doc.add_heading("Frame Captures", level=1)
        for capture in frame_captures:
            key = capture.get("storage_key", "")
            timestamp_sec = capture.get("timestamp_sec", 0.0)
            image_bytes = (frame_bytes_map or {}).get(key)
            if image_bytes:
                try:
                    doc.add_picture(io.BytesIO(image_bytes))
                    doc.add_paragraph(f"Frame @ {timestamp_sec:.1f}s — {key}")
                except Exception:
                    doc.add_paragraph(f"Frame @ {timestamp_sec:.1f}s (image unreadable): {key}")
            else:
                doc.add_paragraph(f"Frame @ {timestamp_sec:.1f}s: {key} (not available)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
